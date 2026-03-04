#!/usr/bin/env python3
"""
generate_angebot.py — Schulter Design Website-Angebot Generator
Max. 2 Seiten. Professioneller Briefcharakter.

Usage:
  python3 generate_angebot.py \
    --firma "Tischlerei Fasching OG" \
    --adresse "Mogersdorf 205, 8382 Mogersdorf" \
    --anrede "Sehr geehrte Herren Fasching," \
    --preview "http://72.60.17.112:8080/" \
    --admin-url "http://72.60.17.112:8080/admin/" \
    --admin-user "admin" \
    --admin-pass "fasching2026" \
    --preis "1.390" \
    --nummer "2026/01" \
    --output "angebot.docx"
"""

import argparse, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Farben ───────────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1B, 0x3A, 0x6B)
CHARCOAL = RGBColor(0x1A, 0x1A, 0x1A)
DARK     = RGBColor(0x2C, 0x2C, 0x2C)
GRAY     = RGBColor(0x55, 0x55, 0x55)
GRAY_LT  = RGBColor(0xAA, 0xAA, 0xAA)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BG_BLUE  = RGBColor(0xF4, 0xF7, 0xFC)   # sehr helles Blau für Vorschau-Block

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH  = os.path.join(SCRIPT_DIR, '..', 'assets', 'schulter-design-logo.jpg')


# ── Helfer ───────────────────────────────────────────────────────────────────

def set_bg(cell, hex6: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6); tcPr.append(shd)


def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'), 'none'); tb.append(el)
    tblPr.append(tb)


def outer_border(tbl, color='D8E0F0', sz='4'):
    tblPr = tbl._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tb.append(el)
    tblPr.append(tb)


def thin_rule(doc, color='D8E4F5', sz='4', before=6, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '1');   b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)
    return p


def section_label(doc, text, before=10, after=4, size=14):
    """Bold navy section heading — clearly visible, generous spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    return p


def cp(cell, text='', bold=False, size=9, color=None,
        align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4, italic=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color or DARK
    return p


# ── Hauptfunktion ─────────────────────────────────────────────────────────────

def build_angebot(firma, adresse, anrede, preview_url, admin_url,
                  admin_user, admin_pass, preis, nummer, datum, output_path):

    doc = Document()

    # Enge Ränder für 2-Seiten-Limit
    for sec in doc.sections:
        sec.top_margin    = Cm(1.6)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════
    # HEADER — Logo | Kontakt
    # ═══════════════════════════════════════════════════════════════
    ht = doc.add_table(rows=1, cols=2)
    no_borders(ht)
    ht.columns[0].width = Cm(8.5)
    ht.columns[1].width = Cm(8.5)

    lc = ht.cell(0, 0)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        lp.add_run().add_picture(LOGO_PATH, width=Cm(5.4))
    else:
        r = lp.add_run('SCHULTER DESIGN')
        r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

    cc = ht.cell(0, 1)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    cp_h = cc.paragraphs[0]
    cp_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cp_h.paragraph_format.space_after = Pt(0)
    for i, (txt, bold, sz, col) in enumerate([
        ('Webdesign  ·  Werbegraafik', True,  8.5, CHARCOAL),
        ('www.schulter-design.at',     False, 8.5, NAVY),
        ('+43 (0) 664 / 45 30 873',   False, 8.5, GRAY),
        ('schulterm@me.com',           False, 8.5, GRAY),
    ]):
        r = cp_h.add_run(('' if i == 0 else '\n') + txt)
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = col

    # Einzige Navy-Trennlinie (einzeln, nicht doppelt)
    thin_rule(doc, color='1B3A6B', sz='6', before=4, after=6)

    # ═══════════════════════════════════════════════════════════════
    # ADRESSBLOCK
    # ═══════════════════════════════════════════════════════════════
    at = doc.add_table(rows=1, cols=2)
    no_borders(at)
    at.columns[0].width = Cm(10)
    at.columns[1].width = Cm(7)

    ac = at.cell(0, 0)
    ap = ac.paragraphs[0]
    ap.paragraph_format.space_before = Pt(10)
    ap.paragraph_format.space_after  = Pt(1)
    for i, line in enumerate([firma] + [l.strip() for l in adresse.split(',')]):
        if i == 0:
            r = ap.add_run(line)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = CHARCOAL
        else:
            p2 = ac.add_paragraph(line)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(1)
            r = p2.runs[0]; r.font.size = Pt(10); r.font.color.rgb = CHARCOAL

    mc = at.cell(0, 1)
    mp = mc.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mp.paragraph_format.space_before = Pt(14)
    r = mp.add_run(f'Mogersdorf, {datum}')
    r.font.size = Pt(9); r.font.color.rgb = GRAY
    np2 = mc.add_paragraph()
    np2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    np2.paragraph_format.space_before = Pt(3)
    rl = np2.add_run('Angebot Nr.  ')
    rl.font.size = Pt(8); rl.font.color.rgb = GRAY
    rn = np2.add_run(nummer)
    rn.bold = True; rn.font.size = Pt(10); rn.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════════
    # BETREFF
    # ═══════════════════════════════════════════════════════════════
    subj = doc.add_paragraph()
    subj.paragraph_format.space_before = Pt(6)
    subj.paragraph_format.space_after  = Pt(14)
    r = subj.add_run(f'Angebot – Professionelle Website für {firma}')
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = CHARCOAL

    # ═══════════════════════════════════════════════════════════════
    # ANREDE + EINLEITUNG
    # ═══════════════════════════════════════════════════════════════
    p = doc.add_paragraph(anrede)
    p.paragraph_format.space_after = Pt(8)
    p.runs[0].font.size = Pt(11); p.runs[0].font.color.rgb = DARK

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    r = intro.add_run(
        'bei einem Besuch Ihrer Website ist mir aufgefallen, dass Sie von einer modernen, '
        'professionellen Online-Präsenz erheblich profitieren könnten. Daher erlaube ich mir, '
        'Ihnen ohne vorherige Anfrage folgendes Angebot zu unterbreiten — ganz unverbindlich.'
    )
    r.font.size = Pt(11); r.font.color.rgb = DARK

    # ═══════════════════════════════════════════════════════════════
    # WARUM EINE WEBSITE — Bullets mit Symbolen
    # ═══════════════════════════════════════════════════════════════
    section_label(doc, 'Warum eine moderne Website entscheidend ist', before=10, after=4)

    bullets = [
        ('🌐', '24 / 7 erreichbar',       'Kunden finden Sie jederzeit – auch außerhalb der Öffnungszeiten.'),
        ('📱', 'Mobil-optimiert',          'Über 70 % aller Suchanfragen kommen heute vom Smartphone.'),
        ('🔍', 'Google-Sichtbarkeit',      'Wer online nicht gefunden wird, verliert Aufträge an die Konkurrenz.'),
        ('✅', 'Vertrauen & Image',         'Ein professioneller Auftritt schafft Vertrauen – noch vor dem ersten Gespräch.'),
        ('📬', 'Direkter Erstkontakt',     'Interessenten können sofort anfragen – ohne langes Suchen.'),
        ('💡', 'Langfristige Investition', 'Einmalige Kosten, dauerhafter Nutzen – der ROI zeigt sich ab dem ersten Auftrag.'),
    ]

    for icon, title, desc in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.3)
        r = p.add_run(f'{icon}  {title}  — {desc}')
        r.font.size = Pt(10); r.font.color.rgb = DARK

    # ═══════════════════════════════════════════════════════════════
    # LEISTUNGSUMFANG — kompakte 2-spaltige Liste
    # ═══════════════════════════════════════════════════════════════
    section_label(doc, 'Leistungsumfang (inkl. Adminbereich)', before=10, after=4)

    features = [
        ('Individuelles Design & Corporate Identity',  'Admin-Panel: Texte & Bilder per Drag & Drop'),
        ('Hero-Slideshow mit Bildwechsel',             'Scroll-Animationen & moderne Übergänge'),
        ('Leistungs-Sektion mit Bildkacheln',          'Mobiloptimierung (Smartphone, Tablet, Desktop)'),
        ('Über uns, Kundenstimmen, FAQ',               'SEO-Grundoptimierung (Meta, Sitemap, Schema)'),
        ('Kontaktformular + Google Maps Verlinkung',   'Social Media Verlinkung im Footer'),
        ('Bürozeiten & Adresse',                       'Impressum & Datenschutz (DSGVO-konform)'),
        ('Bildergalerie',                              'Hosting-Setup (Nginx, SSL-Vorbereitung)'),
    ]

    ft = doc.add_table(rows=len(features), cols=2)
    no_borders(ft)
    ft.columns[0].width = Cm(8.5)
    ft.columns[1].width = Cm(8.5)

    for i, (left, right) in enumerate(features):
        bg = 'F4F7FC' if i % 2 == 0 else 'FFFFFF'
        for col_i, txt in enumerate([left, right]):
            cell = ft.cell(i, col_i)
            set_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.25)
            check = p.add_run('✓  ')
            check.font.size = Pt(8.5); check.font.color.rgb = NAVY; check.bold = True
            text_r = p.add_run(txt)
            text_r.font.size = Pt(8.5); text_r.font.color.rgb = DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ═══════════════════════════════════════════════════════════════
    # LIVE-VORSCHAU — hell, clean, gut lesbar
    # ═══════════════════════════════════════════════════════════════
    section_label(doc, 'Ihre Website – Live-Vorschau & Zugang', before=10, after=5, size=16)

    prev_tbl = doc.add_table(rows=3, cols=2)
    outer_border(prev_tbl, color='C8D8F0', sz='6')
    prev_tbl.columns[0].width = Cm(4.2)
    prev_tbl.columns[1].width = Cm(12.8)

    rows_data = [
        ('🌐  Vorschau-Link',  preview_url,           False),
        ('🔐  Admin-Bereich',  admin_url,             False),
        ('👤  Login',
         f'Benutzername: {admin_user}     Passwort: {admin_pass}', True),
    ]

    for i, (label, value, is_login) in enumerate(rows_data):
        lc = prev_tbl.cell(i, 0)
        vc = prev_tbl.cell(i, 1)
        bg_l = 'EBF0FA' if i % 2 == 0 else 'E4EBF7'
        bg_v = 'F8FAFF' if i % 2 == 0 else 'F2F6FF'
        set_bg(lc, bg_l); set_bg(vc, bg_v)

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after  = Pt(8)
        lp.paragraph_format.right_indent = Cm(0.3)
        lr = lp.add_run(label)
        lr.bold = True; lr.font.size = Pt(8.5); lr.font.color.rgb = NAVY

        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(8)
        vp.paragraph_format.space_after  = Pt(8)
        vp.paragraph_format.left_indent  = Cm(0.4)

        if is_login:
            # Benutzername und Passwort mit klarer Trennung
            parts = value.split('     ')
            for j, part in enumerate(parts):
                k, v = part.split(': ', 1)
                kr = vp.add_run(k + ':  ')
                kr.font.size = Pt(8.5); kr.font.color.rgb = GRAY
                vr = vp.add_run(v)
                vr.bold = True; vr.font.size = Pt(11); vr.font.color.rgb = CHARCOAL
                if j == 0:
                    vp.add_run('          ')
        else:
            vr = vp.add_run(value)
            vr.bold = True; vr.font.size = Pt(10); vr.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════════
    # PREIS
    # ═══════════════════════════════════════════════════════════════
    section_label(doc, 'Angebotssumme', before=8, after=4, size=16)

    pt = doc.add_table(rows=2, cols=3)
    outer_border(pt, color='C8D8F0', sz='6')
    pt.columns[0].width = Cm(1.2)
    pt.columns[1].width = Cm(12.3)
    pt.columns[2].width = Cm(3.5)

    # Header
    for col_i, (txt, align) in enumerate([
        ('Pos.',     WD_ALIGN_PARAGRAPH.CENTER),
        ('Leistung', WD_ALIGN_PARAGRAPH.LEFT),
        ('Betrag',   WD_ALIGN_PARAGRAPH.RIGHT),
    ]):
        c = pt.cell(0, col_i)
        set_bg(c, '1B3A6B')
        cp(c, txt, bold=True, size=8.5, color=WHITE, align=align, sb=5, sa=5)

    # Zeile
    for col_i, (txt, bold, align, col, bg) in enumerate([
        ('1',
         False, WD_ALIGN_PARAGRAPH.CENTER, GRAY,    'FFFFFF'),
        (f'Onepager-Website inkl. Admin-Panel  ·  {firma}',
         True,  WD_ALIGN_PARAGRAPH.LEFT,   CHARCOAL,'FFFFFF'),
        (f'€ {preis},-',
         True,  WD_ALIGN_PARAGRAPH.RIGHT,  NAVY,    'FFFFFF'),
    ]):
        c = pt.cell(1, col_i)
        set_bg(c, bg)
        p_c = c.paragraphs[0]
        p_c.alignment = align
        p_c.paragraph_format.space_before = Pt(7)
        p_c.paragraph_format.space_after  = Pt(7)
        r = p_c.add_run(txt)
        r.bold = bold; r.font.size = Pt(10 if col_i == 2 else 9.5)
        r.font.color.rgb = col

    # Summe
    sum_row = pt.add_row()
    for col_i, (txt, bold, align, col, bg) in enumerate([
        ('',        False, WD_ALIGN_PARAGRAPH.CENTER, GRAY, 'EEF2FA'),
        ('Gesamt',  True,  WD_ALIGN_PARAGRAPH.LEFT,   NAVY, 'EEF2FA'),
        (f'€ {preis},-', True, WD_ALIGN_PARAGRAPH.RIGHT, NAVY, 'EEF2FA'),
    ]):
        c = sum_row.cells[col_i]
        set_bg(c, bg)
        p_c = c.paragraphs[0]
        p_c.alignment = align
        p_c.paragraph_format.space_before = Pt(6)
        p_c.paragraph_format.space_after  = Pt(6)
        r = p_c.add_run(txt)
        r.bold = bold; r.font.size = Pt(11); r.font.color.rgb = col

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ═══════════════════════════════════════════════════════════════
    # ABSCHLUSS
    # ═══════════════════════════════════════════════════════════════
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    nr = note.add_run(
        'Ich besitze keine UID-Nummer, da ich als KleinunternehmerIn umsatzsteuerbefreit bin. '
        'Alle Preise sind Nettopreise. Angebot gültig für 30 Tage. Bei Rückfragen stehe ich '
        'jederzeit gerne zur Verfügung.'
    )
    nr.italic = True; nr.font.size = Pt(8.5); nr.font.color.rgb = GRAY_LT

    for txt, bold, sz, col, sa in [
        ('Mit freundlichen Grüßen,',                     True,  11, DARK,     26),
        ('Mario Schulter, BA BA MA',                      True,  11, CHARCOAL,  3),
        ('Schulter Design  ·  Webdesign & Werbegraafik', False, 10, GRAY,       0),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = col

    # ═══════════════════════════════════════════════════════════════
    # WORD-FUSSZEILE — auf jeder Seite
    # ═══════════════════════════════════════════════════════════════
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after  = Pt(0)

        # Trennlinie über Footer via Border
        fpPr = fp._p.get_or_add_pPr()
        fpBdr = OxmlElement('w:pBdr')
        top_el = OxmlElement('w:top')
        top_el.set(qn('w:val'),   'single')
        top_el.set(qn('w:sz'),    '4')
        top_el.set(qn('w:space'), '2')
        top_el.set(qn('w:color'), 'C8D8F0')
        fpBdr.append(top_el); fpPr.append(fpBdr)

        footer_text = (
            f'Angebot Nr. {nummer}     '
            f'Mario Schulter, BA BA MA  ·  Mogersdorf 253  ·  8382 Mogersdorf  ·  '
            f'Bankverbindung · Kontoinhaber: Mario Schulter · '
            f'Raiffeisen Regionalbank Güssing-Jennersdorf · '
            f'IBAN: AT74 3302 7000 0321 7023 · BIC: RLBBAT2E027'
        )
        # "Angebot Nr. XX" in Navy-Bold, Rest in GRAY_LT
        parts = footer_text.split('     ', 1)
        r_nr = fp.add_run(parts[0])
        r_nr.bold = True; r_nr.font.size = Pt(7); r_nr.font.color.rgb = NAVY
        r_rest = fp.add_run('     ' + parts[1])
        r_rest.font.size = Pt(7); r_rest.font.color.rgb = GRAY_LT

    doc.save(output_path)
    print(f'✅  Gespeichert: {output_path}')


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--firma',       default='Ihr Unternehmen')
    p.add_argument('--adresse',     default='Straße 1, PLZ Ort')
    p.add_argument('--anrede',      default='Sehr geehrte Damen und Herren,')
    p.add_argument('--preview',     default='http://vorschau.at/')
    p.add_argument('--admin-url',   default='http://vorschau.at/admin/')
    p.add_argument('--admin-user',  default='admin')
    p.add_argument('--admin-pass',  default='passwort')
    p.add_argument('--preis',       default='1.390')
    p.add_argument('--nummer',      default=f'2026/{datetime.now().strftime("%m")}')
    p.add_argument('--datum', default=(
        datetime.now().strftime('%d. ') +
        ['Jänner','Februar','März','April','Mai','Juni',
         'Juli','August','September','Oktober','November','Dezember'
        ][datetime.now().month - 1] +
        datetime.now().strftime(' %Y')
    ))
    p.add_argument('--output', default='angebot-website.docx')
    args = p.parse_args()

    build_angebot(
        firma=args.firma, adresse=args.adresse, anrede=args.anrede,
        preview_url=args.preview, admin_url=args.admin_url,
        admin_user=args.admin_user, admin_pass=args.admin_pass,
        preis=args.preis, nummer=args.nummer,
        datum=args.datum, output_path=args.output,
    )

if __name__ == '__main__':
    main()
